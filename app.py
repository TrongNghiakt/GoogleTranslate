from flask import Flask, render_template, request, jsonify, send_file
from googletrans import Translator, LANGUAGES
import os
import pandas as pd
from docx import Document
import fitz 
from io import BytesIO
import time
from deep_translator import GoogleTranslator

app = Flask(__name__)


# Thư mục lưu file đã dịch
TRANSLATION_DIR = 'translations'
if not os.path.exists(TRANSLATION_DIR):
    os.makedirs(TRANSLATION_DIR)

language_codes = list(LANGUAGES.keys())
language_names = list(LANGUAGES.values())


@app.route("/")
def index():
    target_languages = [(code, name) for code, name in LANGUAGES.items()]
    
    print("DEBUG - Target Languages:", target_languages)  # Kiểm tra dữ liệu
    
    return render_template("index.html", source_languages=zip(language_codes, language_names),target_languages=target_languages)

@app.route('/translate', methods=['POST'])
def translate_text():

    text_to_translate = request.form.get('text')
    source_language = request.form.get('source_language')
    target_language = request.form.get('target_language')

    if not text_to_translate or not source_language or not target_language:
        return jsonify({'translation': 'Please provide valid input.'})

    translator = Translator()
    try:
        translation = translator.translate(text_to_translate, src=source_language, dest=target_language).text
    except Exception as e:
        print(f"Translation error: {e}")
        translation = "Translation failed. Please try again later."

    return jsonify({'translation': translation})

@app.route('/pronounce', methods=['POST'])
def pronounce_text():

    text = request.json.get('text')
    if text:
        return jsonify({'message': 'Text pronounced successfully!'})
    else:
        return jsonify({'message': 'Text not provided for pronunciation.'})


@app.route("/translate_file", methods=['POST'])
def translate_file():
    print(request.form)
    target_language = request.form.get('target_language_file')
    print("Selected Language:",target_language)

    if not target_language:
        return render_template('index.html', error="Please select a target language!")
    
    if 'file' not in request.files:
        return render_template('index.html', error="No file uploaded!")


    file = request.files['file']
    target_language = request.form.get("target_language_file")


    filename = file.filename
    if not filename:
        return render_template('index.html', error="No file selected!")
    
    translator = Translator()  # Khởi tạo Translator trước

    try:
        output_path = os.path.join(TRANSLATION_DIR, f'translated_{filename}')
        # Xử lý file TXT
        if filename.endswith('.txt'):
            content = file.read().decode('utf-8')  
            print(f"DEBUG - File Content:\n{content}")
            translated_content = safe_translate(content, target_language)

            
             # Kiểm tra nếu nội dung không thay đổi, không hiển thị link tải
            if translated_content.strip() == content.strip():
                return render_template('index.html', error="Translation failed. Please try again!")

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(translated_content)

        # Xử lý file word (.docx)
        elif filename.endswith('.docx'):
            doc = Document(file)  # Đọc trực tiếp từ file
            translated_doc = Document()
            for para in doc.paragraphs:
                translated_text = safe_translate(para.text, target_language)
                translated_doc.add_paragraph(translated_text)
            
            translated_doc.save(output_path)

        # Xử lý file PDF
        elif filename.endswith('.pdf'):
            doc = fitz.open(stream=file.read(),filetype="pdf")
            text = "\n".join([page.get_text("text", sort=True) for page in doc])

            if not text.strip():  # Kiểm tra nếu file rỗng
                return render_template('index.html', error="PDF file is empty!",
                                       source_languages=zip(language_codes, language_names),
                               target_languages=[(code, name) for code, name in LANGUAGES.items()])
            
            translated_text = safe_translate(text, target_language)


            if translated_text.strip() == text.strip():
                return render_template('index.html', error="Translation failed. Please try again!",
                                       source_languages=zip(language_codes, language_names),
                               target_languages=[(code, name) for code, name in LANGUAGES.items()])

            new_doc = fitz.open()
            page = new_doc.new_page()
            rect = fitz.Rect(50,100,550,800)
            
            # Đường dẫn font
            font_path = "D:/GoogleTranslate/fonts/NotoSansSC-Regular.ttf"


            page.insert_textbox(rect, translated_text, fontsize=12, fontfile=font_path, align=0)


            # Lưu file PDF
            output_path = os.path.join(TRANSLATION_DIR, f'translated_{filename}')
            new_doc.save(output_path)
            new_doc.close()
            

        # Xử lý file Excel
        elif filename.endswith(('.xlsx', '.xls')):
            df = pd.read_excel(file, engine="openpyxl", header=None)

            # Kiểm tra dữ liệu đầu vào
            print("📌 Thông tin file:")
            print(df.info())
            print("📌 Trước khi dịch:", df.head())

            # Xử lý nếu dữ liệu bị rỗng hoặc lỗi
            if df.empty:
                return render_template('index.html', error="⚠️ File Excel không có dữ liệu hoặc bị lỗi!")

            df = df.dropna(how='all', axis=0)  # Xóa dòng rỗng
            df = df.dropna(how='all', axis=1)  # Xóa cột rỗng
            df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)  # Xóa khoảng trắng

            print("📌 Sau khi xử lý dữ liệu rỗng:", df.head())

            def translate_if_text(cell_value, target_language):
                if isinstance(cell_value, str) and any(c.isalpha() for c in cell_value):  # Nếu có chữ cái thì dịch
                    return safe_translate(cell_value, target_language)
                return cell_value  # Nếu là số hoặc trống thì giữ nguyên
            
            df = df.applymap(lambda x: translate_if_text(x, target_language) if pd.notna(x) else x)

            # Kiểm tra dữ liệu sau dịch
            if df.empty:
                return render_template('index.html', error="⚠️ Dữ liệu bị lỗi sau khi dịch!")

            # Lưu file
            output = BytesIO()
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df.to_excel(writer, sheet_name="Translated", index=False,header = False)

            output.seek(0)
            with open(output_path, "wb") as f:
                f.write(output.getvalue())

            print("✅ File dịch đã lưu:", output_path)


        else:
            return render_template('index.html',
                                   source_languages=zip(language_codes, language_names),
                                   target_languages=[(code, name) for code, name in LANGUAGES.items()],
                                   error="Unsupported file format! Only .txt, .docx, .pdf, .xlsx, .xls are allowed.")
        if os.path.exists(output_path):
            return render_template('index.html', 
                       file_result=f'/download/{os.path.basename(output_path)}',
                       source_languages=zip(language_codes, language_names),
                       target_languages=[(code, name) for code, name in LANGUAGES.items()],
                       selected_target_language=target_language)

        return render_template('index.html',source_languages=zip(language_codes, language_names),
                               target_languages=[(code, name) for code, name in LANGUAGES.items()],error="File translation failed!")
                        

    except Exception as e:
        print(f"File translation error: {e}")
        return render_template('index.html',
                               source_languages=zip(language_codes, language_names),
                               target_languages=[(code, name) for code, name in LANGUAGES.items()],
                               error=f"File translation failed: {str(e)}")

@app.route('/download/<file_name>')
def download_file(file_name):
    try:
        file_path = os.path.join(TRANSLATION_DIR, file_name)
        if not os.path.exists(file_path):
            return "File not found!", 404
        response = send_file(file_path, as_attachment=True)
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response
    except Exception as e:
        print(f"Download error: {e}")
        return "File download failed!", 500

def safe_translate(text, target_language, retries=5):
    for i in range(retries):
        try:
            translated_text = GoogleTranslator(source='auto', target=target_language).translate(text)
            if translated_text and translated_text != text:  # Kiểm tra nếu dịch thành công
                return translated_text
        except Exception as e:
            print(f"⚠️ Translation failed (attempt {i+1}): {e}")
            time.sleep(3)  # Tăng thời gian chờ lên 3 giây
    return "⚠️ Translation failed due to connection issues."


if __name__ == '__main__':
    app.run(debug=True, port=4000)
    